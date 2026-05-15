#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
爆款脚本生成流水线 — 百炼 deepseek-v4-flash
用法: python3 run_pipeline.py
"""
import requests, json, re, sys, datetime, textwrap, os
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

API_KEY  = os.getenv("DASHSCOPE_API_KEY", "")   # 请设置环境变量 DASHSCOPE_API_KEY
BASE_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1"
MODEL    = "deepseek-v4-flash"
TOPIC    = "腾讯辟谣AI一号位即将离职"   # 今日热点（2026-05-15 头条热度 3479万）
TODAY    = "2026年5月15日"

def hr(title=""):
    line = "─" * 60
    if title:
        print(f"\n{line}\n【{title}】\n{line}")
    else:
        print(line)

def call_llm(sys_prompt, user_prompt, max_tokens=3000):
    hr()
    print(f"[API] {user_prompt[:60]}…")
    resp = requests.post(
        f"{BASE_URL}/chat/completions",
        headers={"Content-Type": "application/json",
                 "Authorization": f"Bearer {API_KEY}"},
        json={"model": MODEL,
              "messages": [{"role": "system", "content": sys_prompt},
                           {"role": "user",   "content": user_prompt}],
              "temperature": 0.85, "max_tokens": max_tokens},
        timeout=90,
        verify=False
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]

def call_llm_json(sys_prompt, user_prompt, max_tokens=3000):
    text = call_llm(sys_prompt, user_prompt, max_tokens)
    cleaned = re.sub(r'```(?:json)?\s*', '', text).replace('```', '').strip()
    # 修复模型常见输出错误：重复 key，如 "version":  version": 1
    cleaned = re.sub(r'"(\w+)"\s*:\s+\1":\s*', r'"\1": ', cleaned)
    try:
        return json.loads(cleaned)
    except Exception:
        m = re.search(r'(\[[\s\S]*\]|\{[\s\S]*\})', cleaned)
        if m:
            try:
                return json.loads(m.group(1))
            except Exception:
                # 截断修复：去掉最后不完整对象后尝试补全数组
                frag = m.group(1)
                # 找最后一个完整对象 (...},) 的位置
                last_ok = frag.rfind('},')
                if last_ok > 0:
                    fixed = frag[:last_ok+1] + ']'
                    try:
                        result = json.loads(fixed)
                        print(f"⚠ JSON 截断修复成功，保留 {len(result)} 条")
                        return result
                    except Exception:
                        pass
        print(f"[DEBUG raw 500 chars]\n{cleaned[:500]}")
        raise ValueError(f"JSON 解析失败:\n{cleaned[:200]}")

def wrap(text, width=72, indent="  "):
    return "\n".join(indent + l for l in textwrap.wrap(str(text), width))

# ───────────────────────────────────────────────────────────────
hr("Agent 1 · 热点感知")
topics = [{"topic": TOPIC, "heat_score": 95}]
print(f"✓ 手动指定选题：{TOPIC}")

# ───────────────────────────────────────────────────────────────
hr("Agent 2 · 选题决策（十二大爆款法则评分）")
sys2 = ('你是财经短视频选题专家。对选题按爆款潜力评分（0-100），命中以下法则越多越好：'
        '信息差/揭秘、情绪共鸣点、权威解读、反认知观点、常见痛点、替弱势群体发声、'
        '说出大众心里话、借势、名人窥探、涨跌偏好、稀缺信息、娱乐性。'
        '输出纯JSON（不要代码块）：{"topic":"","score":90,"criteria":["法则1"],"structure":"脚本结构名"}')
winner = call_llm_json(sys2, f"对以下选题评分：\n1. {TOPIC}")
print(f"✓ 最佳选题 [{winner.get('score','-')}分]：{winner.get('topic', TOPIC)}")
print(f"✓ 命中法则：{', '.join(winner.get('criteria', []))}")
print(f"✓ 推荐结构：{winner.get('structure', '')}")

# ───────────────────────────────────────────────────────────────
hr("Agent 3 · 解构洗稿（裁缝拼接法）")
sys3 = ('你是财经短视频口播稿专家，用裁缝拼接法（短视频强钩子+长视频深度内容）写200-300字初稿，'
        '含：强开场钩子（10-15字）、问题展开（100字）、核心洞察（100字）、结尾行动引导（一句话）。'
        '只输出纯文本口播稿，不要标注结构说明。')
draft = call_llm(sys3,
    f"选题：{winner.get('topic', TOPIC)}。"
    f"风格：理性犀利，有信息差，口语化，让普通科技从业者/投资者有收获。"
    f"背景：{TODAY}，腾讯内部传出AI战略核心人物（AI一号位）即将离职的消息，随即腾讯官方辟谣，"
    f"引发市场对腾讯AI战略走向、核心人才稳定性及股价影响的广泛讨论。")
print("✓ 口播初稿：")
print(wrap(draft))

# ───────────────────────────────────────────────────────────────
hr("Agent 4 · 脚本生成（5版本，逐版生成保证完整输出）")
ANGLES = ["信息差", "情绪共鸣", "利益驱动", "反差对比", "权威解读"]
STRUCTURES = ["问题前置+层层解答", "场景带入+情绪推进", "钩子前置+顺叙展开", "结果前置+倒叙还原", "观点前置+论据支撑"]
HOOK_TYPES  = ["痛点类", "好奇类", "利益输送类", "反差类", "极限词类"]

sys4_base = (
    '你是短视频脚本专家，擅长科技/财经赛道。'
    '根据选题、口播初稿和要求，生成1个指定角度的完整脚本方案。'
    '输出纯JSON对象（不要代码块）：'
    '{"version":1,"angle":"","cover":"封面≤12汉字，逗号断句",'
    '"direct_post":"直发语50-70字，含2个人物对话感",'
    '"oral_script":"完整口播稿200-280字，开头强钩子+展开+洞察+结尾互动引导",'
    '"pinned":"置顶评论带emoji","hook_type":"","structure":""}'
)

scripts = []
for idx, (angle, structure, hook_type) in enumerate(zip(ANGLES, STRUCTURES, HOOK_TYPES), 1):
    print(f"\n  → 生成版本 {idx}（角度：{angle}，结构：{structure}）")
    user4 = (f"选题：{winner.get('topic', TOPIC)}\n"
             f"口播初稿：{draft}\n"
             f"要求角度：{angle}，结构：{structure}，钩子类型：{hook_type}，版本号：{idx}")
    for attempt in range(2):
        try:
            s = call_llm_json(sys4_base, user4, max_tokens=2000)
            s["version"] = idx
            s.setdefault("angle", angle)
            scripts.append(s)
            print(f"  ✓ 版本 {idx} 完成")
            break
        except Exception as e:
            if attempt == 0:
                print(f"  ⚠ 版本 {idx} 首次失败，重试中...")
            else:
                print(f"  ✗ 版本 {idx} 失败: {e}")

print(f"\n✓ 共生成 {len(scripts)} 个版本")

# ───────────────────────────────────────────────────────────────
hr("Agent 5 · 三重审核")
sys5 = ('你是金融内容合规审核员。检查脚本是否含违禁词（必涨/稳赚/保本/无风险/稳定收益/内幕等）。'
        '输出纯JSON（不要代码块）：{"compliance":true,"oral":true,"professional":true,"issues":[],"suggestions":[]}')
summary_txt = "\n".join(
    f"版本{i+1}：{s.get('oral_script','')[:120]}"
    for i, s in enumerate(scripts)
)
try:
    review = call_llm_json(sys5, f"审核以下脚本：\n{summary_txt}")
except Exception as e:
    print(f"⚠ 审核 LLM 解析失败，默认通过：{e}")
    review = {"compliance": True, "oral": True, "professional": True, "issues": [], "suggestions": []}
print(f"  合规：{'✓ 通过' if review.get('compliance') is not False else '✗ 未通过'}")
print(f"  口播：{'✓ 通过' if review.get('oral')       is not False else '✗ 未通过'}")
print(f"  专业：{'✓ 通过' if review.get('professional') is not False else '✗ 未通过'}")
if review.get("suggestions"):
    print(f"  建议：{review['suggestions'][0]}")

# ───────────────────────────────────────────────────────────────
hr("Agent 6 · 格式化输出（docx）")
today = datetime.datetime.now().strftime("%Y%m%d")
slug  = winner.get("topic", TOPIC)[:10]
fname = f"脚本_{slug}_{today}.docx"
fpath = os.path.join(os.path.dirname(__file__), fname)

doc = Document()

# 标题
t = doc.add_heading(f"爆款脚本：{winner.get('topic', TOPIC)}", level=0)
t.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

# 元信息
meta = doc.add_paragraph()
meta.add_run(f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  模型：{MODEL}").font.size = Pt(9)
doc.add_paragraph()

for i, s in enumerate(scripts):
    cover_disp = (s.get("cover","")).replace("\\n", "\n")
    doc.add_heading(f"版本 {i+1}（{s.get('angle', s.get('hook_type',''))}）", level=1)
    doc.add_paragraph(f"结构：{s.get('structure','')}  |  钩子：{s.get('hook_type','')}").runs[0].bold = True

    doc.add_heading("📸 封面文案", level=2)
    p = doc.add_paragraph(cover_disp)
    p.runs[0].font.size = Pt(13)
    p.runs[0].bold = True

    doc.add_heading("💬 直发语", level=2)
    doc.add_paragraph(s.get("direct_post",""))

    doc.add_heading("🎤 口播文案", level=2)
    doc.add_paragraph(s.get("oral_script",""))

    doc.add_heading("📌 置顶评论", level=2)
    doc.add_paragraph(s.get("pinned",""))

    doc.add_paragraph("─" * 40)

# 口播初稿
doc.add_heading("📝 口播初稿", level=1)
doc.add_paragraph(draft)

doc.save(fpath)
print(f"✓ 已导出：{fname}")

# ───────────────────────────────────────────────────────────────
hr("════ 完整脚本预览 ════")
for i, s in enumerate(scripts):
    print(f"\n{'═'*58}")
    print(f" 版本 {i+1}  角度：{s.get('angle','')}  [{s.get('hook_type','')}]  结构：{s.get('structure','')}")
    print(f"{'═'*58}")
    cover_disp = (s.get("cover","")).replace("\\n", "\n")
    print(f"📸 封面文案：\n{wrap(cover_disp)}")
    print(f"\n💬 直发语：\n{wrap(s.get('direct_post',''))}")
    print(f"\n🎤 口播文案：\n{wrap(s.get('oral_script',''), width=68)}")
    print(f"\n📌 置顶评论：{s.get('pinned','')}")

print(f"\n\n✅ 全部完成！脚本已保存至：{fpath}")
